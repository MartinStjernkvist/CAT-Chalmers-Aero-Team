import re
from pathlib import Path
from docx import Document
from typing import Dict

class ReportParser:
    """Simple parser for progress report documents"""

    @staticmethod
    def read_docx(file_path: Path) -> str:
        """Read all text from a .docx file"""
        try:
            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs]
            
            # Also read text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs.append(cell.text)
            
            return '\n'.join(paragraphs)
        except Exception as e:
            raise Exception(f"Error reading {file_path.name}: {str(e)}")

    @staticmethod
    def split_by_weeks(text: str) -> Dict[int, str]:
        """
        Split document into weeks. Looks for two patterns:
        1. Standalone lines: W42, W.42, w42, w.42
        2. IPR header lines: IPR Name W42, IPR [Name] W42
        """
        weeks = {}
        
        # Pattern 1: Standalone week markers
        standalone_pattern = re.compile(r'^\s*[Ww]\.?\s*(\d+)\s*$', re.MULTILINE)
        
        # Pattern 2: IPR header with week number (IPR ... W42)
        ipr_pattern = re.compile(r'^IPR\s+.+?\s+[Ww]\.?\s*(\d+)', re.MULTILINE)
        
        # Try standalone pattern first
        matches = list(standalone_pattern.finditer(text))
        
        # If no standalone markers, try IPR header pattern
        if not matches:
            matches = list(ipr_pattern.finditer(text))
        
        if not matches:
            return weeks
        
        # Extract content between week markers
        for i, match in enumerate(matches):
            week_num = int(match.group(1))
            start = match.start()
            
            # End is either the start of next week or end of document
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            
            weeks[week_num] = text[start:end].strip()
        
        return weeks

    @staticmethod
    def extract_section(text: str, section_name: str) -> str:
        """Extract a specific section (Progress, Problems, Plans, Documentation)"""
        
        # Define section headers
        sections = {
            'progress': r'^\s*(?:Activities\s*\n\s*)?Progress\s*$',
            'problems': r'^\s*Problems\s*$',
            'plans': r'^\s*Plans\s*$',
            'documentation': r'^\s*Documentation\s*$'
        }
        
        section_pattern = sections.get(section_name.lower())
        if not section_pattern:
            return ""
        
        # Find the start of this section
        match = re.search(section_pattern, text, re.MULTILINE | re.IGNORECASE)
        if not match:
            return ""
        
        start = match.end()
        
        # Find the start of the next section (any section)
        next_section = None
        for name, pattern in sections.items():
            if name == section_name.lower():
                continue
            next_match = re.search(pattern, text[start:], re.MULTILINE | re.IGNORECASE)
            if next_match:
                next_pos = start + next_match.start()
                if next_section is None or next_pos < next_section:
                    next_section = next_pos
        
        # Extract content
        end = next_section if next_section else len(text)
        content = text[start:end].strip()
        
        # Clean up content
        content = re.sub(r'\n{3,}', '\n\n', content)
        
        # Remove subsection labels
        content = re.sub(r'^\s*(General|Completed tasks|Specific tasks)\s*$', '', content, flags=re.MULTILINE)
        
        # Remove IPR question prompts
        content = re.sub(r'^\s*What problems am I encountering\?\s*$', '', content, flags=re.MULTILINE)
        content = re.sub(r"^\s*What['’]s your plan moving forward\?\s*$", '', content, flags=re.MULTILINE)
        content = re.sub(r'^\s*What have I achieved this week\?\s*$', '', content, flags=re.MULTILINE)
        
        # Remove SPR question prompts
        content = re.sub(r'^\s*What has the sub-team achieved this week\?\s*$', '', content, flags=re.MULTILINE)
        content = re.sub(r'^\s*What problems are we encountering\?\s*$', '', content, flags=re.MULTILINE)
        content = re.sub(r"^\s*What['’]s the plan moving forward\?\s*$", '', content, flags=re.MULTILINE)
        
        # Remove other common template prompts
        content = re.sub(r'^\s*Have I learnt any valuable lessons\?\s*$', '', content, flags=re.MULTILINE)
        content = re.sub(r'^\s*Youtube / Literature / Websites / Documents\s*$', '', content, flags=re.MULTILINE)
        content = re.sub(r'^\s*Links to documents or contributions to reports\s*$', '', content, flags=re.MULTILINE)
        content = re.sub(r'^\s*Learning resources\s*$', '', content, flags=re.MULTILINE)
        content = re.sub(r'^\s*Lessons learnt\s*$', '', content, flags=re.MULTILINE)
        content = re.sub(r'^\s*Documented material\s*$', '', content, flags=re.MULTILINE)
        
        # Clean up excessive newlines again
        content = re.sub(r'\n{3,}', '\n\n', content)
        
        return content.strip()

    @staticmethod
    def extract_metadata(text: str) -> Dict[str, str]:
        """Extract metadata like sub-team and main task"""
        metadata = {}
        
        # Extract sub-team
        sub_team_match = re.search(
            r'Sub-team\s*\(current\)\s*:\s*(.+?)(?:\n|$)',
            text,
            re.IGNORECASE
        )
        if sub_team_match:
            metadata['sub_team'] = sub_team_match.group(1).strip()
        
        # Extract main task
        task_match = re.search(
            r'Main task\s*\(current\)\s*:\s*(.+?)(?:\n|$)',
            text,
            re.IGNORECASE
        )
        if task_match:
            metadata['main_task'] = task_match.group(1).strip()
        
        return metadata